# encoding=utf-8
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import time
from cn2an import an2cn

from lib.numConvert import float2int
from exam_system.settings import TEACHING_FTP_ROOT


def genDocx(dict_data: dict, exam_manage):
    # dict_data 是考试试卷数据
    # exam_manage 是 django model ORM的实例数据

    # 解析考试相关数据
    exam_name = exam_manage.exam_name
    exam_type = exam_manage.exam_type
    start_time = exam_manage.start_time
    end_time = exam_manage.end_time
    duration = exam_manage.duration
    max_score = exam_manage.max_score
    multiple_half = exam_manage.multiple_half

    # 创建word文档
    document = Document()

    # 设置中文字体
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 封皮
    title_echo = document.add_heading(level=0)
    title_run = title_echo.add_run('{exam_name}{exam_type}试卷'.format(exam_name=exam_name, exam_type=exam_type))
    # 设置字体样式
    # title_run.bold = True  # 设置为黑体
    title_run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
    title_run.font.name = '宋体'  # 设置为宋体

    p = document.add_paragraph()
    # 输出考试时间
    exam_data_echo = start_time.strftime('%Y-%m-%d')
    start_time_echo = start_time.strftime('%H:%M')
    end_time_echo = end_time.strftime('%H:%M')
    # p.add_run("考试时间： ").bold = True
    p.add_run("考试时间：")
    exam_time_echo = '{exam_data_echo} {start_time_echo}-{end_time_echo}'.format(exam_data_echo=exam_data_echo, start_time_echo=start_time_echo, end_time_echo=end_time_echo)
    p.add_run(exam_time_echo)

    # 添加换行
    p.add_run().add_break()

    # 输出卷面总分
    p.add_run("总分："+str(float2int(max_score)))

    # 统计 单选、多选、填空、简答、论述 题数
    statics_question_types = {
        '单选': 0,
        '多选': 0,
        '填空': 0,
        '简答': 0,
        '论述': 0,
    }
    for question in dict_data:
        if question['question_type']=='单选':
            statics_question_types['单选'] += 1
        if question['question_type']=='多选':
            statics_question_types['多选'] += 1
        if question['question_type']=='填空':
            statics_question_types['填空'] += 1
        if question['question_type']=='简答':
            statics_question_types['简答'] += 1
        if question['question_type']=='论述':
            statics_question_types['论述'] += 1

    current_question_type_flag = None
    current_question_type_num = 1
    question_number_set = 1

    for question in dict_data:

        # 不同题型的判分框
        if question['question_type'] != current_question_type_flag:


            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "总分"
            table.width = Inches(3)
            table.style = 'Table Grid'
            for row in table.rows:
                for cell in row.cells:
                    cell.width = Inches(0.5)

            p = document.add_paragraph()
            cn_number = an2cn(current_question_type_num)


            if question['question_type']=="单选":
                current_question_type_flag = "单选"

                p.add_run("{cn_number}、{question_type}题".format(
                    cn_number=cn_number,
                    question_type=current_question_type_flag,
                ))

                if statics_question_types['单选']<=10:
                    table = document.add_table(rows=2, cols=statics_question_types['单选']+1)
                    table.style = 'Table Grid'
                    table.cell(0, 0).text = "题号"
                    table.cell(1, 0).text = "答案"
                    for i in range(1,statics_question_types['单选']+1):
                        table.cell(0, i).text = str(i)
                else:
                    lines = int(statics_question_types['单选']/10)    # 完整的行数
                    last_line_items = int(statics_question_types['单选']%10)  # 完整行数外还剩几个
                    if last_line_items==0:
                        rows = 2*lines
                    else:
                        rows = 2 * (lines+1)
                    table = document.add_table(rows=rows, cols=11)
                    table.style = 'Table Grid'
                    for text_row_header in range(0, int(rows/2)):
                        table.cell(text_row_header*2, 0).text = "题号"
                        table.cell(text_row_header*2+1, 0).text = "答案"
                    for text_row_header in range(0, int(rows/2)):
                        for i in range(1,11):
                            table.cell(text_row_header * 2, i).text = str(i+text_row_header*10)

            elif question['question_type']=="多选":
                current_question_type_flag = "多选"

                if multiple_half:
                    multi_select_additional_des = "少选得一半分"
                else:
                    multi_select_additional_des = "少选不得分"

                p.add_run("{cn_number}、{question_type}题（{multi_select_additional_des}）".format(
                    cn_number=cn_number,
                    question_type=current_question_type_flag,
                    multi_select_additional_des=multi_select_additional_des,
                ))

                if statics_question_types['多选']<=10:
                    table = document.add_table(rows=2, cols=statics_question_types['多选']+1)
                    table.style = 'Table Grid'
                    table.cell(0, 0).text = "题号"
                    table.cell(1, 0).text = "答案"
                    for i in range(1,statics_question_types['多选']+1):
                        table.cell(0, i).text = str(i)
                else:
                    lines = int(statics_question_types['多选']/10)    # 完整的行数
                    last_line_items = int(statics_question_types['多选']%10)  # 完整行数外还剩几个
                    if last_line_items==0:
                        rows = 2*lines
                    else:
                        rows = 2 * (lines+1)
                    table = document.add_table(rows=rows, cols=11)
                    table.style = 'Table Grid'
                    for text_row_header in range(0, int(rows/2)):
                        table.cell(text_row_header*2, 0).text = "题号"
                        table.cell(text_row_header*2+1, 0).text = "答案"
                    for text_row_header in range(0, int(rows/2)):
                        for i in range(1,11):
                            table.cell(text_row_header * 2, i).text = str(i+text_row_header*10)
            elif question['question_type'] == "填空":
                current_question_type_flag = "填空"
                p.add_run("{cn_number}、{question_type}题".format(
                    cn_number=cn_number,
                    question_type=current_question_type_flag,
                ))
            elif question['question_type'] == "简答":
                current_question_type_flag = "简答"
                p.add_run("{cn_number}、{question_type}题".format(
                    cn_number=cn_number,
                    question_type=current_question_type_flag,
                ))
            elif question['question_type'] == "论述":
                current_question_type_flag = "论述"
                p.add_run("{cn_number}、{question_type}题".format(
                    cn_number=cn_number,
                    question_type=current_question_type_flag,
                ))
            current_question_type_num+=1
        p = document.add_paragraph()

        # 输出题号、问题
        p.add_run("{question_number_set}. {question_text}（{score}分）".format(
            question_number_set=question_number_set,
            question_text=question['question_text'],
            score=float2int(question['score']),
        ))

        # 选项设置左侧缩进
        # p = document.add_paragraph()
        # p_format = p.paragraph_format # 设置段落格式
        # p_format.left_indent = Pt(20)  # 左侧缩进
        # p_format.first_line_indent = Pt(-20)  # 首行缩进

        # 选择题，含单选、多选
        # 输出选项、回答空行
        if question['question_type'] in ('单选', "多选"):
            p = document.add_paragraph()
            p_format = p.paragraph_format
            p_format.left_indent = Pt(20)

            # 建立ABCD索引
            options = question['choice_text'].split("|")
            options_dict = dict()
            for i, option in enumerate(options):
                options_dict[chr(65 + i)] = option
            for index, (key, single_option) in enumerate(options_dict.items()):
                option_echo = "{index}、{single_option}".format(index=key, single_option=single_option)
                if index != len(options_dict) - 1:
                    p.add_run(option_echo)
                    p.add_run().add_break()
                else:
                    p.add_run(option_echo)
            # 取消左侧缩进设置
            # p = document.add_paragraph()
            # p.paragraph_format.space_before = 0
            # p.paragraph_format.space_after = 0

        elif question['question_type'] == "填空":
            p.add_run().add_break()

        elif question['question_type'] == "简答":
            # 添加多个空白段落（行）
            for _ in range(2):  # 假设要添加3行空白
                p = document.add_paragraph()

        elif question['question_type'] == "论述":
            # 添加多个空白段落（行）
            for _ in range(4):  # 假设要添加3行空白
                p = document.add_paragraph()

        # 题号自增
        question_number_set = question_number_set + 1
    filepath = os.path.join(TEACHING_FTP_ROOT, '{exam_name}.docx'.format(exam_name=exam_name))
    # print(filepath)
    document.save(filepath)
    return '{exam_name}.docx'.format(exam_name=exam_name)
